# -*- coding: utf-8 -*-
#!/usr/bin/python
import requests, Config, os, sys, time, urllib, urllib2, shutil, zipfile,  unicodedata #docx,
import ftplib, time
from subprocess import Popen, PIPE
from docx import Document
from docx.shared import Inches, RGBColor, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import socket

from flask import stream_with_context, Response
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle, StyleSheet1
from reportlab.lib.units import inch, mm
from reportlab.lib.pagesizes import letter, inch, landscape
from reportlab.lib import colors

####################################################
# Possivel erro que pode dar: a imagem logo do PDF. 
# Comente a setenca da imagem do logo
####################################################

def CreateDir(NameDir):
	#print '\n'
	if not os.path.exists(Config.ProjectDir):
		try:
			os.mkdir(Config.ProjectDir)
			os.chdir(Config.ProjectDir) 
		except Exception as e:
			#print e
			return u'Falha de criação de diretório. Verifique permissões da pasta'	

	pathDirectory = Config.ProjectDir + NameDir
	if not os.path.exists(pathDirectory):
		try:
			os.mkdir(pathDirectory)
			os.chdir(pathDirectory) 
			os.mkdir('Fotos')
			os.mkdir('Audios')
			os.chdir(pathDirectory+'/Audios') 
		except Exception as e:
			#print e
			return u'Falha de criação de diretório. Verifique permissões da pasta'


def transferAudios(file, i, NameDir):	
	#print '\n'
	#print '****************************'
	#print '* DOWNLOAD AUDIO '+str(i)+'*'
	#print '****************************'		
			
	if file['path']:
		try:
			shutil.copyfile(file['path'], Config.ProjectDir+NameDir+'/Audios/'+str(file['cod'])+"_gsm.wav")
		except Exception as e:
			print 'ERRO1:',e, '\nDiretorio ou arquivo existente no BD e inexistente no FTP'
			return True
	else:
		print 'VAZIO.',	file['path']
	return True



###########################  ANTIGO ###########################################
"""def transferAudios(file, i):	
	#print '\n'
	#print '****************************'
	#print '* DOWNLOAD AUDIO '+str(i)+'*'
	#print '****************************'		
	try:		
		ftp = ftplib.FTP(Config.RemoteIP, Config.Username, Config.Password) 		
		#print file
		if file['path']:
			oldArchive = file['path'].replace("\\","/")
			path= oldArchive.rsplit('/', 1)
			try:
				ftp.cwd('/'+path[0]) 
				newArchive = open(str(file['cod'])+"_gsm.wav", 'wb')
				ftp.retrbinary('RETR %s' % path[1], newArchive.write)
				newArchive.close()
			except Exception as e:
				print e
				print 'Diretorio ou arquivo existente no BD e inexistente no FTP'
				return True
		else:
			print 'VAZIO.',	file['path']
		ftp.close()	
		return True

	except (ftplib.error_temp, IOError) as f:
		print f,' ', str(i)
		time.sleep(2)
		transferAudios(file,i)

	except Exception as e:	
		print 'ERRO:',e	
		return False"""
###########################  ANTIGO ###########################################


def nameDir(name):
	if name:
		for i in name:
			if ' ' in i['nom']:
				nome= i['nom'].replace(" ", "_")
				NameDir= ''.join((c for c in unicodedata.normalize('NFD', nome) if unicodedata.category(c) != 'Mn'))+'_' 
			else: 
				NameDir= ''.join((c for c in unicodedata.normalize('NFD', i['nom']) if unicodedata.category(c) != 'Mn'))+'_'
	else:
		NameDir= 'TodasGravacoes_'
	data = time.strftime("%d_%m_%Y-%H.%M.%S")

	return NameDir+data

def transferPhotos(item, NameDir, i): 
	#print '*****************************'
	#print '*  DOWNLOAD FOTO '+str(i)+' *'
	#print '*****************************'

	URLphotos = Config.ProjectDir+NameDir+'/Fotos/'
	if item['fot1']:
		URLget = Config.Photos+item['fot1']
		if not os.path.exists(URLphotos+item['fot1']):
			try:
				response = urllib2.urlopen(URLget, URLphotos+item['fot1']) 
				response.close()
			except Exception as e:
				print e
			else:
				urllib.urlretrieve(URLget, URLphotos+item['fot1']) 	
			
def convertAudio(file, nameDir, i):
	#print '*********************************'
	#print '* CONVERSAO AUDIO SOX '+str(i)+'*'
	#print '*********************************'

	nameDir = Config.ProjectDir+nameDir+'/Audios/'
	path= nameDir+str(file['cod'])+'_gsm.wav'
	newPath= nameDir+str(file['cod'])+'.wav'
	DirPath = Config.ProjectDir.rsplit('/',2)[0]

	cmd = '"'+DirPath+'/sox" "'+path+'" -e signed-integer "'+newPath+'"'

	if os.path.isfile(path):
		try:	
			os.chdir(nameDir)
			process = Popen(cmd, shell=True, stdout=PIPE, stderr=PIPE)
			stdout, stderr = process.communicate()
		except Exception as e:
			print e
		else:
			try:
				os.remove(path)			
			except Exception as e:
				print  e
	
def staticFiles(nameDir):
	path= Config.ProjectDir+nameDir
	pathDirectory= path+'/static'
	staticDir = Config.ProjectDir.rsplit('/',2)[0]

	if not os.path.exists(pathDirectory):
		try:
			os.mkdir(pathDirectory)
			shutil.copy(staticDir+'/static/favicon.ico', pathDirectory)
			shutil.copy(staticDir+'/static/style.css', pathDirectory)
			return True
		except Exception as e:
			print e
			return False

def HTML(resp, nameDir):
	path= Config.ProjectDir+nameDir
	os.chdir(path)
	Html_file= open("WebTigerDownloads.html","w")
	try:
		resp=resp.encode(encoding='UTF-8',errors="ignore")
		Html_file.write(resp)
		return True
	except Exception as e:
			print e
			return False
	Html_file.close()

def PDF(nameDir, fileRows, callswithPhoto):
	try:
		os.chdir(Config.ProjectDir+nameDir)
		doc = SimpleDocTemplate('WebTigerDownloads.pdf', pagesize=letter, topMargin=15, bottomMargin=25, rightMargin=10,leftMargin=10) #, showBoundary=True
		doc.title ='WebTigerDownloads.pdf'

		styleSheet = getSampleStyleSheet()
		styleSheet2 = StyleSheet1()
		styleSheet2.add(ParagraphStyle(
        name = 'BodyText2',
        parent = styleSheet['BodyText'],
        borderColor = '#000000',
        borderWidth = 1,
        borderPadding = (7, 7, 7, 7)))

		logo = Config.ProjectDir.rsplit('/',2)[0]+'/static/logo.png'
		#<img src='+logo+' height="20" width="20" valign="middle"/> 
		P0 = Paragraph('<para align=center spaceb=3> <font color="green">WEBTIGER DOWNLOADS</font></para>', styleSheet["Heading1"])
		elements = []
		P1 = Paragraph('<b>Total de Registros: '+str(len(fileRows))+'</b>', styleSheet["BodyText"])

		elements.append(P0)	
		elements.append(P1)		

		for index, item in enumerate(fileRows):
			info= StringCell(fileRows, index)
			elements.append (Paragraph('<br/><br/>', styleSheet["BodyText"]))
			if callswithPhoto!='semFoto':
				if item['fot1']:
					I = Image(item['fot1'], hAlign = 'CENTER')
					I.drawHeight = 1.00*inch*I.drawHeight / I.drawWidth
					I.drawWidth = 1.00*inch
					elements.append(I)
				elements.append(Paragraph(StringPDF(info), styleSheet['BodyText'])) #styleSheet2['BodyText2']
			else:
				elements.append(Paragraph(StringPDF(info), styleSheet2['BodyText2']))

		doc.build(elements, onFirstPage=addPageNumber, onLaterPages=addPageNumber)
		return True

	except Exception as e:
		print e		
		return False

def StringPDF(info):
	String= '<b>ID Chamada: </b>'+ str(info['ID Chamada'])+'<br/><b>Data:</b> '+str(info['Data'])+ u'<br/><b>Duração:</b> '+str(info['Duracao'])+ u'<br/><b>Direção:</b> '+str(info['Direcao']) + '<br/><b>Alvo:</b> '+info['Alvo']
	String= String + '<br/><b>Telefone:</b> '+info['Telefone'] +'<br/><b>IMEI:</b> '+info['IMEI'] 
	String = String +'<br/><b>Interlocutor:</b> '+info['Interlocutor']+u'<br/><b>Transcrição:</b> '+info['Transcricao']+'<br/><br/><b>Sinopse:</b> '+info['Sinopse']+u'<br/><br/><b>Áudio:</b> ' 
	if info['Path']:
		audio = info["Path"].rsplit('/',1)[1]
		String = String + '<a href="file:///Audios/'+audio+'" color="blue">'+audio+'</a><br/>' 
	elif info['Memo']:
		String = String + '<b><font color="red">'+info['Memo']+'</font></b>'
	else:
		String = String + '<b><font color="red">Audio inexistente no Banco de Dados ou no servidor FTP</font></b>'
	return String

def addPageNumber(canvas, doc):
    #Add the page number
    page_num = canvas.getPageNumber()
    text = "%s" % page_num
    canvas.drawRightString(205*mm, 5*mm, text)

def WORD(nameDir, fileRows, callswithPhoto):
	try:
		os.chdir(Config.ProjectDir+nameDir)
		document = docx.Document()
		logo= Config.ProjectDir.rsplit('/',2)[0]+'/static/logo.png'

		sections = document.sections
		for section in sections:
		    section.top_margin = Cm(1.0)
		    section.bottom_margin = Cm(1.0)
		    section.left_margin = Cm(1.0)
		    section.right_margin = Cm(1.0)

		#======= CABECALHO
		p1 = document.add_paragraph()
		p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
		pp1= p1.add_run()
		pp1.bold= True
		pp1.font.size = Pt(14)
		pp1.font.name = 'Arial'
		pp1.font.color.rgb = RGBColor(0, 76, 25)
		pp1.add_picture(logo, width=Inches(0.3))
		pp1.add_text(' WEBTIGER DOWNLOADS')

		#======= TOTAL REGISTROS 
		p2 = document.add_paragraph().add_run()
		p2.add_text('Total de registros: '+str(len(fileRows))) 
		p2.font.size = Pt(12)
		p2.font.name = 'Arial'
		p2.bold= True

		#======= TABELA
		for index, item in enumerate(fileRows):
			if callswithPhoto!='semFoto':
				aux=document.add_paragraph('\n')
				aux.alignment = WD_ALIGN_PARAGRAPH.CENTER
				foto= aux.add_run()
				if item['fot1']:
					foto.add_picture(item['fot1'], width=Inches(0.6))

			table = document.add_table(rows=3, cols=2, style="TableGrid") 
			insercaoWord(document, table, StringCell(fileRows, index))
		document.save('WebTigerDownloads.docx')
		return True
	except Exception as e:
		print e
		return False	


def insercaoWord(document, table, data):
	P1= table.cell(0, 0).add_paragraph()
	P11 = P1.add_run('Alvo: ')
	P12= P1.add_run(data['Alvo'])
	P11.font.size = Pt(11)
	P11.font.name = 'Arial'
	P11.bold= True
	P12.font.size = Pt(11)
	P12.font.name = 'Arial'

	P2= table.cell(0, 1).add_paragraph()
	P21 = P2.add_run('Telefone do Alvo: ')
	P22= P2.add_run(data['Telefone'])
	P21.font.size = Pt(11)
	P21.font.name = 'Arial'
	P21.bold= True
	P22.font.size = Pt(11)
	P22.font.name = 'Arial'

	P3= table.cell(1, 0).add_paragraph()
	P31 = P3.add_run('Data: ')
	P32= P3.add_run(data['Data'])
	P31.font.size = Pt(11)
	P31.font.name = 'Arial'
	P31.bold= True
	P32.font.size = Pt(11)
	P32.font.name = 'Arial'

	P4= table.cell(1, 1).add_paragraph()
	P41 = P4.add_run(u'Duração: ')
	P42= P4.add_run(data['Duracao'])
	P41.font.size = Pt(11)
	P41.font.name = 'Arial'
	P41.bold= True
	P42.font.size = Pt(11)
	P42.font.name = 'Arial'

	P5= table.cell(2, 0).add_paragraph()
	P51 = P5.add_run('Registro: ') #ou nomeInterlocutor
	if data['Path']:
		registro = data['Path'].rsplit('/',3)
		P52= P5.add_run(registro[1]+'/'+registro[2]+'/'+registro[3]) 
	else:
		P52= P5.add_run(data['Path']) 
	P51.font.size = Pt(11)
	P51.font.name = 'Arial'
	P51.bold= True
	P52.font.size = Pt(11)
	P52.font.name = 'Arial'

	P6= table.cell(2, 1).add_paragraph()
	P61 = P6.add_run('Telefone Interlocutor: ')
	P62= P6.add_run(data['Interlocutor'])
	P61.font.size = Pt(11)
	P61.font.name = 'Arial'
	P61.bold= True
	P62.font.size = Pt(11)
	P62.font.name = 'Arial'

	P7= document.add_paragraph('\n')
	P71= P7.add_run(u'Transcrição: ')  
	P71.font.size = Pt(11)
	P71.font.name = 'Arial'
	P71.bold= True
	P72= P7.add_run(data['Transcricao'])
	P72.font.size = Pt(11)
	P72.font.name = 'Arial'

	P8= document.add_paragraph('\n')
	P81 = P8.add_run('Sinopse: ')
	P82= P8.add_run(data['Sinopse'])
	P81.font.size = Pt(11)
	P81.font.name = 'Arial'
	P81.bold= True
	P82.font.size = Pt(11)
	P82.font.name = 'Arial'

	P9= document.add_paragraph()
	P91 = P9.add_run(u'Clique para escutar o áudio: ')
	if data['Path']: 
		path = data['Path'].rsplit('/',1)[1]
		add_hyperlink(P9, './Audios/'+path, path)
	elif data['Memo']:
		P92 = P9.add_run(data['Memo'])
		P92.font.color.rgb = RGBColor(255, 0, 0)
	else:
		P92 = P9.add_run('Audio inexistente no Banco de Dados ou no servidor FTP')
		P92.font.color.rgb = RGBColor(255, 0, 0)
	P91.font.size = Pt(11)
	P91.font.name = 'Arial'
	P91.bold= True

	P10= document.add_paragraph()
	P101 = P10.add_run(u'Segure Ctrl e de um clique encima do nome do áudio para tocá-lo')	
	P101.font.size = Pt(8)
	P101.font.name = 'Arial'
	P101.font.color.rgb = RGBColor(0, 76, 25)
	P101.bold= True

def StringCell(fileRows, index):
	if not fileRows[index]['nom']:
		nome= ''
	else:
		nome = fileRows[index]['nom']
		try:
			nome= fileRows[index]['nom'].decode(encoding='UTF-8',errors='strict')
		except Exception as e:
			print e	
			try:
				nome= fileRows[index]['nom'].decode(encoding='latin-1',errors='strict') 
			except Exception as e:
				print e

	if not fileRows[index]['trs']:
		trans= ''
	else: 
		trans = fileRows[index]['trs'] 
		try:
			trans= fileRows[index]['trs'].decode(encoding='UTF-8',errors='strict')
		except Exception as e:
			print e
			try:
				trans= fileRows[index]['trs'].decode(encoding='latin-1',errors='strict') 
			except Exception as e:
				print e

	if not fileRows[index]['obs']:
		obs= ''
	else: 
		obs = fileRows[index]['obs']
		try:
			obs= fileRows[index]['obs'].decode(encoding='UTF-8',errors='strict')
		except Exception as e:
			print e
			try:
				obs= fileRows[index]['obs'].decode(encoding='latin-1',errors='strict') 
			except Exception as e:
				print e				

	if not fileRows[index]['memo']:
		memo= ''
	else: 
		memo = fileRows[index]['memo']
		try:
			memo= fileRows[index]['memo'].decode(encoding='UTF-8',errors='strict')
		except Exception as e:
			print e
			try:
				memo= fileRows[index]['memo'].decode(encoding='latin-1',errors='strict') 
			except Exception as e:
				print e						
			
	if not fileRows[index]['apl']:
		apelido= ''
	else: 
		apelido = fileRows[index]['apl']
		try:
			apelido	= fileRows[index]['apl'].decode(encoding='UTF-8',errors='strict')
		except Exception as e:
			print e
			try:
				apelido	= fileRows[index]['apl'].decode(encoding='latin-1',errors='strict') 
			except Exception as e:
				print e	

	if not fileRows[index]['nome']:
		interlocutor= ''
	else: 
		interlocutor = fileRows[index]['nome']
		try:
			interlocutor= fileRows[index]['nome'].decode(encoding='UTF-8',errors='strict')
		except Exception as e:
			print e
			try:
				interlocutor= fileRows[index]['nome'].decode(encoding='latin-1',errors='strict') 
			except Exception as e:
				print e					

	if not fileRows[index]['direcao']:
		direcao= ''
	else: 
		direcao= fileRows[index]['direcao']

	if not fileRows[index]['dur']:
		duracao= ''
	else: 
		duracao= str(fileRows[index]['dur'])

	if not fileRows[index]['num']:
		num= ''
	else: 
		num= str(fileRows[index]['num'])
	
	if not fileRows[index]['ddd']:
		ddd= ''
	else: 
		ddd= '('+str(fileRows[index]['ddd'])+')'	

	if not fileRows[index]['imei']:
		imei= ''
	else: 
		imei= str(fileRows[index]['imei'])	

	if not fileRows[index]['num_lig']:
		num_lig= ''
	else: 
		num_lig= str(fileRows[index]['num_lig'])	

	result= {
		'ID Chamada' : str(fileRows[index]['cod']),
		'Data' : str(fileRows[index]['dat_hor']),
		'Duracao' : duracao,
		'Direcao' : direcao,
		'Alvo' : nome+' ('+apelido+')',
		'Telefone' : ddd+num,
		'IMEI' : imei,
		'Interlocutor' : num_lig+' ('+interlocutor+')',
		'Transcricao' : trans,
		'Sinopse' : obs,
		'Memo' : memo, 
		'Path': fileRows[index]['path'],
		'NomeInterlocutor':interlocutor		
	}
		
	return result
	

def copyPDFWord(nameDir, fileRows, callswithPhoto):
	#print '\n'
	#print '*****************************'
	#print '*    COPIANDO ARQUIVOS      *'
	#print '*****************************'

	try:
		os.chdir(Config.ProjectDir+nameDir)
		if os.path.exists(nameDir):	
			shutil.rmtree(nameDir)
		os.mkdir(nameDir)

		os.chdir(Config.ProjectDir+nameDir+'/'+nameDir)
		os.mkdir('Audios')
		if callswithPhoto =='comFoto':
			os.mkdir('Fotos')
		
		shutil.copytree(Config.ProjectDir+nameDir+'/static', Config.ProjectDir+nameDir+'/'+nameDir+'/static')
		shutil.copyfile(Config.ProjectDir+nameDir+'/WebTigerDownloads.html', Config.ProjectDir+nameDir+'/'+nameDir+'/WebTigerDownloads.html')
		if os.path.exists(Config.ProjectDir+nameDir+'/WebTigerDownloads.docx'):
			shutil.copyfile(Config.ProjectDir+nameDir+'/WebTigerDownloads.docx', Config.ProjectDir+nameDir+'/'+nameDir+'/WebTigerDownloads.docx')
		if os.path.exists(Config.ProjectDir+nameDir+'/WebTigerDownloads.pdf'):
			shutil.copyfile(Config.ProjectDir+nameDir+'/WebTigerDownloads.pdf', Config.ProjectDir+nameDir+'/'+nameDir+'/WebTigerDownloads.pdf')
		return True
	except Exception as e:
		print e
		return False

def copyImageAudio(nameDir, item):
	#print item['path'],' ', item['cod'], item['fot1']

	if item['path'] and item['path']!='':
		nameAudio= item['path'].rsplit('/',1)[1]
		#print nameAudio
			
		if os.path.exists(item['path']):
			shutil.copyfile(item['path'], Config.ProjectDir+nameDir+'/'+nameDir+'/Audios/'+nameAudio)

	if item['fot1'] and item['fot1']!='':
		nameFoto= item['fot1'].rsplit('/',1)[1]	
		if (not(os.path.exists(Config.ProjectDir+nameDir+'/'+nameDir+'/Fotos/'+nameFoto))) and (os.path.exists(item['fot1'])):
			shutil.copyfile(item['fot1'], Config.ProjectDir+nameDir+'/'+nameDir+'/Fotos/'+nameFoto)

def zip(nameDir):		
		os.chdir(Config.ProjectDir+nameDir)
		shutil.make_archive(nameDir, 'zip', nameDir)

def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink
